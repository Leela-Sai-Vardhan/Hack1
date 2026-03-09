import os
import logging
from datetime import datetime
from google import genai
from google.genai import types
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import config
from schemas import CreditCase

logger = logging.getLogger(__name__)
_client = None

def _get_client():
    global _client
    if _client is None:
        _client = genai.Client(api_key=config.GEMINI_API_KEY)
    return _client

_FIVE_CS = [
    ("CHARACTER",   "Promoter & Management Quality",
     "promoter background, DIN checks, news sentiment, litigation, management quality"),
    ("CAPACITY",    "Financial Performance & Cash Flows",
     "revenue trends, EBITDA margins, DSCR, interest coverage, GST-Bank reconciliation"),
    ("CAPITAL",     "Balance Sheet Strength",
     "debt-to-equity, tangible net worth, current ratio, related party exposure"),
    ("COLLATERAL",  "Security Coverage",
     "collateral type, coverage ratio, MCA charge registry, encumbrance"),
    ("CONDITIONS",  "Sectoral & Macro Context",
     "sector outlook, RBI/SEBI regulatory risks, peer benchmarking"),
]


def _cell_bg(cell, hex_color: str):
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" '
                    f'w:color="auto" w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shd)


def _gemini_text(prompt: str) -> str:
    try:
        resp = _get_client().models.generate_content(
            model=config.GEMINI_MODEL,
            contents=prompt,
            config=types.GenerateContentConfig(temperature=0.4))
        return resp.text.strip()
    except Exception as e:
        logger.error(f"Gemini CAM generation error: {e}")
        return "Section pending manual review."


def _section_data(key: str, case: CreditCase) -> str:
    m  = case.financial_metrics
    rf = case.research_findings
    pi = case.primary_insights
    if key == "CHARACTER":
        return (f"Promoter integrity: {rf.promoter_integrity_score}/100 | "
                f"Flags: {rf.promoter_flags or 'None'} | News risk: {rf.news_risk_level} | "
                f"{rf.news_summary} | Litigation: {len(rf.litigation_cases)} cases, "
                f"₹{rf.litigation_total_exposure_lakhs:.0f}L | "
                f"Mgmt quality: {pi.management_response_quality or 'not assessed'} | "
                f"Succession plan: {pi.succession_plan_exists}")
    elif key == "CAPACITY":
        return (f"Revenue (Yr1/Yr2/Yr3 ₹L): {m.revenue_yr1}/{m.revenue_yr2}/{m.revenue_yr3} | "
                f"EBITDA margin: {m.ebitda_margin_pct}% | CAGR: {m.revenue_cagr_3yr}% | "
                f"DSCR: {m.dscr} | Int coverage: {m.interest_coverage} | "
                f"GST/Bank ratio: {m.gst_bank_ratio} | "
                f"Capacity observed vs reported: {pi.capacity_utilization_observed_pct}% vs "
                f"{pi.capacity_utilization_reported_pct}% | "
                f"Primary insight adj: {pi.ai_score_adjustment}pts — {pi.ai_adjustment_reason}")
    elif key == "CAPITAL":
        return (f"Total debt ₹{m.total_debt}L | Equity ₹{m.total_equity}L | "
                f"D/E: {m.debt_to_equity} | Current ratio: {m.current_ratio} | "
                f"OD utilisation: {m.od_utilization_pct}")
    elif key == "COLLATERAL":
        return "Collateral details to be added by credit officer post physical verification."
    elif key == "CONDITIONS":
        return (f"Sector: {rf.sector} | Outlook: {rf.sector_outlook} | "
                f"Regulatory risks: {rf.regulatory_risks or 'None identified'}")
    return "Data pending."


def _section_score(key: str, sb) -> int:
    return {
        "CHARACTER":  int(sb.research_score * 0.70 + sb.primary_insight_score * 0.30),
        "CAPACITY":   int(sb.financial_score),
        "CAPITAL":    int(sb.financial_score * 0.85),
        "COLLATERAL": 70,
        "CONDITIONS": int(sb.research_score * 0.50 + 50),
    }.get(key, 65)


def generate_cam(case: CreditCase) -> CreditCase:
    doc = Document()

    # ── Title ─────────────────────────────────────────────────────────────────
    t = doc.add_heading("CREDIT APPRAISAL MEMO (CAM)", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(
        f"Company: {case.company_name}  |  CIN: {case.company_cin or 'N/A'}  |  "
        f"Date: {datetime.now().strftime('%d %b %Y')}  |  Prepared by: Intelli-Credit AI"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ── Executive Summary ─────────────────────────────────────────────────────
    doc.add_heading("EXECUTIVE SUMMARY", 1)
    d  = case.decision
    sb = case.score_breakdown
    rows = [
        ("Risk Rating",       sb.risk_rating),
        ("Composite Score",   f"{sb.final_score}/100"),
        ("Recommendation",    d.recommendation),
        ("Recommended Limit", f"₹{d.recommended_limit_cr} Cr" if d.recommended_limit_cr else "N/A"),
        ("Interest Rate",     f"MCLR + {d.mclr_spread_bps}bps ({d.effective_rate_pct}%)" if d.mclr_spread_bps else "N/A"),
        ("Financial Score",   f"{sb.financial_score}/100"),
        ("Research Score",    f"{sb.research_score}/100"),
        ("Primary Score",     f"{sb.primary_insight_score}/100"),
    ]
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = "Table Grid"
    for i, (lbl, val) in enumerate(rows):
        tbl.cell(i, 0).text = lbl
        tbl.cell(i, 1).text = str(val)
        tbl.cell(i, 0).paragraphs[0].runs[0].bold = True
        if lbl == "Recommendation":
            _cell_bg(tbl.cell(i, 1),
                     {"APPROVE":"C6EFCE","CONDITIONAL":"FFEB9C","DECLINE":"FFC7CE"}.get(val,"FFFFFF"))
    doc.add_paragraph()

    # ── Five Cs ───────────────────────────────────────────────────────────────
    for cs_key, cs_title, focus in _FIVE_CS:
        doc.add_heading(f"{cs_key} — {cs_title}", 1)
        sec_score = min(100, max(0, _section_score(cs_key, sb)))
        p = doc.add_paragraph(f"Section Score: {sec_score}/100")
        p.runs[0].bold = True
        data_ctx = _section_data(cs_key, case)
        narrative = _gemini_text(
            f"Write the {cs_key} section for a Credit Appraisal Memo.\n"
            f"Company: {case.company_name}\n"
            f"Focus areas: {focus}\n"
            f"Data:\n{data_ctx}\n\n"
            f"3 paragraphs, formal Indian banking language, reference specific numbers, "
            f"state risks directly. Use MCLR, NPA, GSTR, DSCR, NCLT where relevant. "
            f"No headers, no JSON — body text only."
        )
        doc.add_paragraph(narrative)
        doc.add_paragraph()

    # ── Risk Flags ────────────────────────────────────────────────────────────
    doc.add_heading("RISK FLAGS (Auto-Detected)", 1)
    if case.risk_flags:
        ft = doc.add_table(rows=1 + len(case.risk_flags), cols=3)
        ft.style = "Table Grid"
        for i, h in enumerate(["Severity", "Source", "Description"]):
            ft.cell(0, i).text = h
            ft.cell(0, i).paragraphs[0].runs[0].bold = True
        for i, flag in enumerate(case.risk_flags, 1):
            ft.cell(i, 0).text = flag.severity
            ft.cell(i, 1).text = flag.source
            ft.cell(i, 2).text = flag.description
            _cell_bg(ft.cell(i, 0),
                     {"HIGH":"FFC7CE","MEDIUM":"FFEB9C","LOW":"C6EFCE"}.get(flag.severity,"FFFFFF"))
    else:
        doc.add_paragraph("No significant risk flags detected.")
    doc.add_paragraph()

    # ── Score Breakdown ───────────────────────────────────────────────────────
    doc.add_heading("SCORE BREAKDOWN & EXPLAINABILITY", 1)
    expl_p = doc.add_paragraph()
    expl_p.add_run("Top Score Drivers:").bold = True
    for drv in sb.score_drivers[:6]:
        sym = "▲" if drv["direction"] == "positive" else ("▼" if drv["direction"] == "negative" else "►")
        doc.add_paragraph(
            f"  {sym} {drv['description']}: {drv['value']}  "
            f"(impact: {drv['impact']:+.1f} pts | sub-score: {drv['sub_score']}/100)"
        )
    doc.add_paragraph()
    rp = doc.add_paragraph()
    rp.add_run("Financial Ratios Used:").bold = True
    for k, v in sb.financial_ratios_used.items():
        doc.add_paragraph(f"  {k}: {v}")
    doc.add_paragraph()

    # ── Final Recommendation ──────────────────────────────────────────────────
    doc.add_heading("FINAL RECOMMENDATION", 1)
    rr = doc.add_paragraph()
    run = rr.add_run(f"DECISION: {d.recommendation}")
    run.bold = True
    run.font.size = Pt(14)

    if d.recommendation == "APPROVE":
        doc.add_paragraph(f"Approved limit: ₹{d.recommended_limit_cr} Cr "
                          f"at {d.effective_rate_pct}% p.a. (MCLR + {d.mclr_spread_bps}bps)")
    elif d.recommendation == "CONDITIONAL":
        doc.add_paragraph(f"Conditional approval for ₹{d.recommended_limit_cr} Cr subject to:")
        for cond in d.key_conditions:
            doc.add_paragraph(f"  • {cond}")
    else:
        doc.add_paragraph("Application declined due to:")
        for reason in d.rejection_reasons:
            doc.add_paragraph(f"  • {reason}")

    doc.add_paragraph()
    ep = doc.add_paragraph()
    ep.add_run("Explainability & Counterfactual:").bold = True
    doc.add_paragraph(sb.counterfactual or sb.decision_explanation)

    # ── Footer ────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    foot = doc.add_paragraph(
        f"Generated by Intelli-Credit AI  |  "
        f"{datetime.now().strftime('%d %b %Y %H:%M')}  |  "
        f"Case: {case.case_id}  |  FOR INTERNAL USE ONLY"
    )
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Save ──────────────────────────────────────────────────────────────────
    safe = "".join(c if c.isalnum() or c in "-_" else "_" for c in case.company_name)
    path = os.path.join(config.OUTPUT_DIR, f"CAM_{safe}_{case.case_id[:8]}.docx")
    doc.save(path)
    logger.info(f"CAM saved: {path}")
    case.cam_path = path
    return case
